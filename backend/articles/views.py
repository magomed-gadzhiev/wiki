from rest_framework import viewsets, status, permissions
from rest_framework.exceptions import ValidationError
from rest_framework.decorators import action
from rest_framework.response import Response
from django.db.models import Q, F
from django.db.models.functions import Lower
from django.utils import timezone
from .models import Article, ArticleVersion, ArticleImage, ArticleAttachment, Category, Model, Technology, Tag, ArticleOption, ArticleOptionValue, Group, ArticleTemplate, Comment
from .serializers import ArticleSerializer, ArticleListSerializer, ArticleVersionSerializer, ArticleImageSerializer, ArticleAttachmentSerializer, CategorySerializer, ModelSerializer, TechnologySerializer, TagSerializer, ArticleOptionSerializer, ArticleOptionValueSerializer, GroupSerializer, GroupDetailSerializer, ArticleTemplateSerializer, CommentSerializer
from .permissions import ArticlePermission
from .word_import_processor import WordImportProcessor
import mammoth
import tempfile
import os
import base64
import zipfile
from bs4 import BeautifulSoup
try:
    import pypandoc
    PYPANDOC_AVAILABLE = True
except ImportError:
    PYPANDOC_AVAILABLE = False
try:
    from docx import Document
    from docx.shared import RGBColor
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


class ArticleViewSet(viewsets.ModelViewSet):
    """
    ViewSet для управления статьями
    """
    queryset = Article.objects.all()
    permission_classes = [ArticlePermission]
    
    def _extract_images_from_docx(self, docx_path):
        """Извлекает изображения из DOCX файла и возвращает словарь {имя_файла: base64_data}"""
        images = {}
        try:
            with zipfile.ZipFile(docx_path, 'r') as docx_zip:
                # Изображения обычно находятся в word/media/
                for file_info in docx_zip.filelist:
                    if file_info.filename.startswith('word/media/'):
                        image_name = os.path.basename(file_info.filename)
                        image_data = docx_zip.read(file_info.filename)
                        
                        # Определяем MIME тип по расширению
                        ext = os.path.splitext(image_name)[1].lower()
                        mime_types = {
                            '.jpg': 'image/jpeg',
                            '.jpeg': 'image/jpeg',
                            '.png': 'image/png',
                            '.gif': 'image/gif',
                            '.bmp': 'image/bmp',
                            '.webp': 'image/webp',
                            '.svg': 'image/svg+xml'
                        }
                        mime_type = mime_types.get(ext, 'image/png')
                        
                        # Конвертируем в base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        images[image_name] = {
                            'data': image_base64,
                            'mime_type': mime_type
                        }
        except Exception as e:
            pass  # Если не удалось извлечь, просто возвращаем пустой словарь
        return images
    
    def _convert_images_to_base64(self, html_content, docx_path):
        """Обрабатывает HTML и заменяет пути к изображениям на base64 data URI"""
        # Извлекаем изображения из DOCX
        images = self._extract_images_from_docx(docx_path)
        
        if not images:
            return html_content
        
        # Парсим HTML
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Находим все теги img
        for img in soup.find_all('img'):
            src = img.get('src', '')
            if not src:
                continue
            
            # Извлекаем имя файла из пути (может быть media/image1.jpeg или просто image1.jpeg)
            image_name = os.path.basename(src)
            
            # Убираем параметры запроса, если есть (например, ?width=100)
            if '?' in image_name:
                image_name = image_name.split('?')[0]
            
            # Ищем изображение в словаре
            matched = False
            if image_name in images:
                image_info = images[image_name]
                img['src'] = f"data:{image_info['mime_type']};base64,{image_info['data']}"
                matched = True
            else:
                # Пробуем найти по части имени (на случай если путь отличается)
                # Сначала пробуем точное совпадение без расширения
                image_name_no_ext = os.path.splitext(image_name)[0]
                for img_name, img_info in images.items():
                    img_name_no_ext = os.path.splitext(img_name)[0]
                    if image_name_no_ext.lower() == img_name_no_ext.lower():
                        img['src'] = f"data:{img_info['mime_type']};base64,{img_info['data']}"
                        matched = True
                        break
                
                # Если не нашли, пробуем частичное совпадение
                if not matched:
                    for img_name, img_info in images.items():
                        if image_name.lower() in img_name.lower() or img_name.lower() in image_name.lower():
                            img['src'] = f"data:{img_info['mime_type']};base64,{img_info['data']}"
                            matched = True
                            break
        
        return str(soup)
    
    def _apply_word_colors_to_html(self, html_content, docx_path):
        """
        Извлекает цвета текста из Word документа и применяет их к HTML.
        """
        if not PYTHON_DOCX_AVAILABLE:
            return html_content
        
        try:
            # Открываем Word документ
            doc = Document(docx_path)
            
            # Парсим HTML
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Извлекаем все текстовые элементы из Word с их стилями (цвет, размер, шрифт)
            # Создаем карту: текст -> стили
            text_style_map = {}
            
            # Обрабатываем параграфы
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    text = run.text.strip()
                    if text and len(text) > 0:
                        styles = {}
                        
                        # Извлекаем цвет
                        if run.font.color and run.font.color.rgb:
                            # Конвертируем RGB в hex
                            rgb = run.font.color.rgb
                            if rgb is not None:
                                try:
                                    rgb_int = None
                                    color_hex = None
                                    
                                    # RGBColor в python-docx может быть представлен по-разному
                                    if isinstance(rgb, RGBColor):
                                        # Пробуем разные способы извлечения int значения
                                        if hasattr(rgb, 'val'):
                                            rgb_int = rgb.val
                                        elif hasattr(rgb, '__int__'):
                                            rgb_int = int(rgb)
                                        else:
                                            # Пробуем получить через строковое представление
                                            rgb_str = str(rgb)
                                            import re
                                            # Может быть в формате RGB(r, g, b)
                                            match = re.search(r'(\d+),\s*(\d+),\s*(\d+)', rgb_str)
                                            if match:
                                                r, g, b = map(int, match.groups())
                                                color_hex = f"#{r:02x}{g:02x}{b:02x}"
                                            else:
                                                # Пробуем извлечь как hex строку
                                                hex_match = re.search(r'[0-9A-Fa-f]{6}', rgb_str)
                                                if hex_match:
                                                    color_hex = f"#{hex_match.group(0)}"
                                    elif isinstance(rgb, int):
                                        rgb_int = rgb
                                    
                                    # Если получили int, извлекаем компоненты
                                    if rgb_int is not None:
                                        r = (rgb_int >> 16) & 0xFF
                                        g = (rgb_int >> 8) & 0xFF
                                        b = rgb_int & 0xFF
                                        color_hex = f"#{r:02x}{g:02x}{b:02x}"
                                    
                                    # Если получили цвет, сохраняем его
                                    if color_hex:
                                        styles['color'] = color_hex
                                except (ValueError, AttributeError, TypeError):
                                    # Если не удалось извлечь цвет, пропускаем
                                    pass
                        
                        # Извлекаем размер шрифта
                        if run.font.size:
                            try:
                                # Размер в пунктах (Points), конвертируем в пиксели или оставляем как есть
                                size_pt = run.font.size.pt
                                styles['font-size'] = f"{size_pt}pt"
                            except (AttributeError, ValueError):
                                pass
                        
                        # Извлекаем семейство шрифта
                        if run.font.name:
                            try:
                                font_name = run.font.name
                                styles['font-family'] = font_name
                            except (AttributeError, ValueError):
                                pass
                        
                        # Сохраняем стили для этого текста
                        if styles:
                            normalized_text = ' '.join(text.split())
                            if normalized_text not in text_style_map:
                                text_style_map[normalized_text] = {}
                            text_style_map[normalized_text].update(styles)
            
            # Обрабатываем таблицы
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                text = run.text.strip()
                                if text and len(text) > 0:
                                    styles = {}
                                    
                                    # Извлекаем цвет
                                    if run.font.color and run.font.color.rgb:
                                        rgb = run.font.color.rgb
                                        if rgb is not None:
                                            try:
                                                rgb_int = None
                                                color_hex = None
                                                
                                                if isinstance(rgb, RGBColor):
                                                    if hasattr(rgb, 'val'):
                                                        rgb_int = rgb.val
                                                    elif hasattr(rgb, '__int__'):
                                                        rgb_int = int(rgb)
                                                    else:
                                                        rgb_str = str(rgb)
                                                        import re
                                                        match = re.search(r'(\d+),\s*(\d+),\s*(\d+)', rgb_str)
                                                        if match:
                                                            r, g, b = map(int, match.groups())
                                                            color_hex = f"#{r:02x}{g:02x}{b:02x}"
                                                        else:
                                                            hex_match = re.search(r'[0-9A-Fa-f]{6}', rgb_str)
                                                            if hex_match:
                                                                color_hex = f"#{hex_match.group(0)}"
                                                elif isinstance(rgb, int):
                                                    rgb_int = rgb
                                                
                                                if rgb_int is not None:
                                                    r = (rgb_int >> 16) & 0xFF
                                                    g = (rgb_int >> 8) & 0xFF
                                                    b = rgb_int & 0xFF
                                                    color_hex = f"#{r:02x}{g:02x}{b:02x}"
                                                
                                                if color_hex:
                                                    styles['color'] = color_hex
                                            except (ValueError, AttributeError, TypeError):
                                                pass
                                    
                                    # Извлекаем размер шрифта
                                    if run.font.size:
                                        try:
                                            size_pt = run.font.size.pt
                                            styles['font-size'] = f"{size_pt}pt"
                                        except (AttributeError, ValueError):
                                            pass
                                    
                                    # Извлекаем семейство шрифта
                                    if run.font.name:
                                        try:
                                            font_name = run.font.name
                                            styles['font-family'] = font_name
                                        except (AttributeError, ValueError):
                                            pass
                                    
                                    # Сохраняем стили для этого текста
                                    if styles:
                                        normalized_text = ' '.join(text.split())
                                        if normalized_text not in text_style_map:
                                            text_style_map[normalized_text] = {}
                                        text_style_map[normalized_text].update(styles)
            
            # Применяем стили к HTML элементам
            # Обрабатываем все текстовые элементы, включая элементы внутри таблиц
            for elem in soup.find_all(['p', 'span', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6', 
                                      'strong', 'em', 'b', 'i', 'u', 'td', 'th', 'li', 'a']):
                # Для ячеек таблиц обрабатываем также вложенные элементы
                if elem.name in ['td', 'th']:
                    # Обрабатываем вложенные элементы в ячейках
                    for nested_elem in elem.find_all(['span', 'p', 'strong', 'em', 'b', 'i', 'u']):
                        self._apply_styles_to_element(nested_elem, text_style_map)
                    # Обрабатываем саму ячейку
                    self._apply_styles_to_element(elem, text_style_map)
                else:
                    self._apply_styles_to_element(elem, text_style_map)
            
            return str(soup)
        except Exception as e:
            # Если не удалось применить стили, возвращаем исходный HTML
            import traceback
            print(f"Ошибка при применении стилей из Word: {e}")
            print(traceback.format_exc())
            return html_content
    
    def _apply_styles_to_element(self, elem, text_style_map):
        """
        Применяет стили (цвет, размер, шрифт) к элементу из карты стилей Word.
        """
        elem_text = elem.get_text().strip()
        normalized_elem_text = ' '.join(elem_text.split()) if elem_text else ''
        
        # Получаем текущие стили элемента
        current_style = elem.get('style', '')
        style_dict = {}
        if current_style:
            # Парсим существующие стили
            parts = [p.strip() for p in current_style.split(';') if p.strip()]
            for part in parts:
                if ':' in part:
                    key, value = part.split(':', 1)
                    style_dict[key.strip().lower()] = value.strip()
        
        # Проверяем, есть ли стили для этого текста
        if normalized_elem_text in text_style_map:
            styles = text_style_map[normalized_elem_text]
            # Применяем стили, если их еще нет
            for style_key, style_value in styles.items():
                if style_key not in style_dict:
                    style_dict[style_key] = style_value
        else:
            # Проверяем частичное совпадение
            for text, styles in text_style_map.items():
                if text in normalized_elem_text and len(text) > 3:
                    # Применяем стили, если их еще нет
                    for style_key, style_value in styles.items():
                        if style_key not in style_dict:
                            style_dict[style_key] = style_value
        
        # Сохраняем обновленные стили
        if style_dict:
            style_string = '; '.join([f'{k}: {v}' for k, v in style_dict.items()])
            elem['style'] = style_string
    
    def get_serializer_class(self):
        if self.action == 'list':
            return ArticleListSerializer
        return ArticleSerializer
    
    def get_queryset(self):
        user = self.request.user
        queryset = Article.objects.all()
        
        # Фильтрация удаленных статей
        # Суперпользователи видят все статьи, включая удаленные
        if not user.is_superuser:
            # Обычные пользователи не видят удаленные статьи
            queryset = queryset.filter(is_deleted=False)

        # Список (list): object-level permission в DRF для list не применяется к queryset,
        # поэтому для уровня read ограничиваем выборку здесь (опубликованные + свои черновики).
        if user.is_authenticated and user.is_active and not user.is_superuser:
            checker = ArticlePermission()
            eff = checker._get_effective_system_permission(user)
            if eff == 'read':
                queryset = queryset.filter(Q(is_published=True) | Q(author=user))
            elif eff == 'none':
                queryset = queryset.none()
        
        # Поиск
        search = self.request.query_params.get('search', None)
        if search:
            queryset = queryset.filter(
                Q(model_name__icontains=search) |
                Q(content__icontains=search) |
                Q(summary__icontains=search)
            )
        
        # Фильтр по автору
        author = self.request.query_params.get('author', None)
        if author:
            queryset = queryset.filter(author_id=author)
        
        # Фильтр по публикации
        is_published = self.request.query_params.get('is_published', None)
        if is_published is not None:
            queryset = queryset.filter(is_published=is_published.lower() == 'true')
        
        # Фильтр по модели
        model_id = self.request.query_params.get('model', None)
        if model_id:
            queryset = queryset.filter(model_id=model_id)
        
        # Фильтр по тегам (можно передать несколько через tags=id1&tags=id2)
        tag_ids = self.request.query_params.getlist('tags')
        if tag_ids:
            queryset = queryset.filter(tags__id__in=tag_ids).distinct()
        
        # Фильтр по опциям статей (можно передать несколько через option_id[] и option_value[])
        option_ids = self.request.query_params.getlist('option_id')
        option_values = self.request.query_params.getlist('option_value')
        if option_ids and option_values and len(option_ids) == len(option_values):
            # Применяем фильтры по опциям (AND между разными опциями)
            import re
            for option_id, option_value in zip(option_ids, option_values):
                if option_id and option_value:
                    # Регистронезависимый поиск по значению опции (используем __iregex для SQLite с кириллицей)
                    # Экранируем специальные символы и добавляем .* для поиска подстроки
                    option_value_escaped = re.escape(option_value)
                    queryset = queryset.filter(
                        option_values__option_id=option_id,
                        option_values__value__iregex=f'.*{option_value_escaped}.*'
                    ).distinct()
        
        return queryset.select_related('author', 'model').prefetch_related(
            'can_view', 'can_edit', 'can_delete', 'tags', 'images', 'versions',
            'option_values__option'
        )
    
    def retrieve(self, request, *args, **kwargs):
        """Увеличиваем счетчик просмотров"""
        instance = self.get_object()
        Article.objects.filter(pk=instance.pk).update(view_count=F('view_count') + 1)
        instance.refresh_from_db()
        serializer = self.get_serializer(instance)
        return Response(serializer.data)
    
    def perform_create(self, serializer):
        """Создание статьи с автоматическим созданием первой версии"""
        article = serializer.save()
        # Создаем первую версию
        ArticleVersion.objects.create(
            article=article,
            model_name=article.model_name,
            content=article.content,
            summary=article.summary,
            version_number=1,
            author=article.author,
            change_description='Создание статьи'
        )
    
    def perform_update(self, serializer):
        """Обновление статьи с созданием новой версии"""
        article = serializer.instance
        old_content = article.content
        old_model_name = article.model_name
        
        # Получаем change_description из request.data, если есть
        change_description = self.request.data.get('change_description', 'Обновление статьи')
        
        # Сохраняем изменения
        updated_article = serializer.save()
        
        # Создаем новую версию, если контент изменился
        if old_content != updated_article.content or old_model_name != updated_article.model_name:
            last_version = article.versions.first()
            new_version_number = (last_version.version_number + 1) if last_version else 1
            
            ArticleVersion.objects.create(
                article=updated_article,
                model_name=updated_article.model_name,
                content=updated_article.content,
                summary=updated_article.summary,
                version_number=new_version_number,
                author=self.request.user,
                change_description=change_description
            )
    
    def destroy(self, request, *args, **kwargs):
        """Мягкое удаление статьи - ставим метку удаления вместо физического удаления"""
        instance = self.get_object()
        from django.utils import timezone
        instance.is_deleted = True
        instance.deleted_at = timezone.now()
        instance.save()
        return Response(status=status.HTTP_204_NO_CONTENT)
    
    @action(detail=True, methods=['post'])
    def publish(self, request, pk=None):
        """Публикация статьи (изменение статуса is_published на True)"""
        article = self.get_object()
        
        if article.is_published:
            return Response(
                {'error': 'Статья уже опубликована'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        article.is_published = True
        article.save()
        
        serializer = self.get_serializer(article)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def unpublish(self, request, pk=None):
        """Снятие статьи с публикации (изменение статуса is_published на False)"""
        article = self.get_object()
        
        if not article.is_published:
            return Response(
                {'error': 'Статья не опубликована'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        article.is_published = False
        article.save()
        
        serializer = self.get_serializer(article)
        return Response(serializer.data)
    
    @action(detail=True, methods=['post'])
    def restore(self, request, pk=None):
        """Восстановление удаленной статьи (только для суперпользователей)"""
        article = self.get_object()
        
        # Проверяем, что пользователь - суперпользователь
        if not request.user.is_superuser:
            return Response(
                {'error': 'Только суперпользователи могут восстанавливать удаленные статьи'}, 
                status=status.HTTP_403_FORBIDDEN
            )
        
        if not article.is_deleted:
            return Response(
                {'error': 'Статья не удалена'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        article.is_deleted = False
        article.deleted_at = None
        article.save()
        
        serializer = self.get_serializer(article)
        return Response(serializer.data)
    
    @action(detail=True, methods=['get'])
    def versions(self, request, pk=None):
        """Получить все версии статьи (только для пользователей с правами редактирования)"""
        article = self.get_object()
        
        # Проверяем права на просмотр версий
        # Суперпользователи и авторы имеют доступ
        if request.user.is_superuser or article.author == request.user:
            versions = article.versions.all()
            serializer = ArticleVersionSerializer(versions, many=True)
            return Response(serializer.data)
        
        # Проверяем системные права через группы
        from .permissions import ArticlePermission
        permission_checker = ArticlePermission()
        system_permission = permission_checker._get_effective_system_permission(request.user)
        
        # Только пользователи с правами редактирования (edit) могут просматривать версии
        if system_permission == 'edit':
            versions = article.versions.all()
            serializer = ArticleVersionSerializer(versions, many=True)
            return Response(serializer.data)
        else:
            # При праве только на чтение или отсутствии прав - запрещаем просмотр версий
            return Response(
                {'error': 'У вас нет прав на просмотр истории версий. Доступно только чтение статей.'}, 
                status=status.HTTP_403_FORBIDDEN
            )
    
    @action(detail=True, methods=['post'])
    def restore_version(self, request, pk=None):
        """Восстановить статью из версии"""
        article = self.get_object()
        version_id = request.data.get('version_id')
        
        if not version_id:
            return Response(
                {'error': 'version_id обязателен'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        try:
            version = article.versions.get(id=version_id)
            article.model_name = version.model_name
            article.content = version.content
            article.summary = version.summary
            article.save()
            
            # Создаем новую версию с пометкой о восстановлении
            last_version = article.versions.first()
            new_version_number = (last_version.version_number + 1) if last_version else 1
            
            ArticleVersion.objects.create(
                article=article,
                model_name=article.model_name,
                content=article.content,
                summary=article.summary,
                version_number=new_version_number,
                author=request.user,
                change_description=f'Восстановление из версии {version.version_number}'
            )
            
            serializer = self.get_serializer(article)
            return Response(serializer.data)
        except ArticleVersion.DoesNotExist:
            return Response(
                {'error': 'Версия не найдена'}, 
                status=status.HTTP_404_NOT_FOUND
            )
    
    @action(detail=True, methods=['post'])
    def upload_image(self, request, pk=None):
        """Загрузить изображение для статьи"""
        article = self.get_object()
        image_file = request.FILES.get('image')
        alt_text = request.data.get('alt_text', '')
        
        if not image_file:
            return Response(
                {'error': 'Изображение обязательно'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        article_image = ArticleImage.objects.create(
            article=article,
            image=image_file,
            alt_text=alt_text,
            uploaded_by=request.user
        )
        
        serializer = ArticleImageSerializer(article_image, context={'request': request})
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @action(detail=True, methods=['post'])
    def upload_attachment(self, request, pk=None):
        """Загрузить файл (вложение) для статьи"""
        article = self.get_object()
        file = request.FILES.get('file')
        comment = request.data.get('comment', '')
        
        # Максимальный размер файла: 50 МБ
        MAX_FILE_SIZE = 50 * 1024 * 1024  # 50 МБ в байтах
        
        if not file:
            return Response(
                {'error': 'Файл обязателен'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        # Проверка размера файла
        if file.size > MAX_FILE_SIZE:
            return Response(
                {'error': f'Размер файла превышает максимально допустимый размер 50 МБ. Размер файла: {file.size / (1024 * 1024):.2f} МБ'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        attachment = ArticleAttachment.objects.create(
            article=article,
            file=file,
            filename=file.name,
            file_size=file.size,
            comment=comment,
            uploaded_by=request.user
        )
        
        serializer = ArticleAttachmentSerializer(attachment, context={'request': request})
        return Response(serializer.data, status=status.HTTP_201_CREATED)
    
    @action(detail=False, methods=['post'])
    def import_word(self, request):
        """Импорт содержимого из Word документа с помощью Pandoc (fallback на mammoth)"""
        word_file = request.FILES.get('file')
        
        if not word_file:
            return Response(
                {'error': 'Файл обязателен'}, 
                status=status.HTTP_400_BAD_REQUEST
            )
        
        converter_used = None
        warnings = []
        
        # Пробуем использовать Pandoc
        if PYPANDOC_AVAILABLE:
            try:
                word_file.seek(0)
                # Сохраняем файл во временный файл для Pandoc
                with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
                    for chunk in word_file.chunks():
                        tmp_file.write(chunk)
                    tmp_file_path = tmp_file.name
                
                try:
                    # Конвертируем с помощью Pandoc с улучшенными параметрами для сохранения стилей
                    # Используем --to html5 для лучшей поддержки стилей
                    html_content = pypandoc.convert_file(
                        tmp_file_path,
                        'html5',
                        format='docx',
                        extra_args=[
                            '--standalone',
                            '--wrap=none',
                            '--preserve-tabs',
                            '--syntax-highlighting=none'  # Отключаем подсветку кода для сохранения оригинальных стилей
                        ]
                    )
                    
                    # Обрабатываем изображения и конвертируем в base64
                    html_content = self._convert_images_to_base64(html_content, tmp_file_path)
                    
                    # Извлекаем цвета из Word документа и применяем к HTML
                    if PYTHON_DOCX_AVAILABLE:
                        html_content = self._apply_word_colors_to_html(html_content, tmp_file_path)
                    
                    converter_used = 'pandoc'
                finally:
                    # Удаляем временный файл
                    if os.path.exists(tmp_file_path):
                        os.unlink(tmp_file_path)
                
            except Exception as pandoc_error:
                # Если Pandoc не сработал, пробуем mammoth
                warnings.append(f'Pandoc не смог обработать файл: {str(pandoc_error)}. Используется fallback на mammoth.')
                converter_used = None
        
        # Fallback на mammoth, если Pandoc не доступен или не сработал
        if not converter_used:
            try:
                word_file.seek(0)
                import base64
                
                # Настраиваем mammoth для извлечения изображений
                def convert_image(image):
                    """Конвертирует изображение в base64"""
                    with image.open() as image_bytes:
                        image_data = image_bytes.read()
                        # Определяем MIME тип
                        content_type = image.content_type
                        if 'jpeg' in content_type or 'jpg' in content_type:
                            mime_type = 'image/jpeg'
                        elif 'png' in content_type:
                            mime_type = 'image/png'
                        elif 'gif' in content_type:
                            mime_type = 'image/gif'
                        elif 'webp' in content_type:
                            mime_type = 'image/webp'
                        else:
                            mime_type = 'image/png'
                        
                        # Конвертируем в base64
                        image_base64 = base64.b64encode(image_data).decode('utf-8')
                        return {
                            "src": f"data:{mime_type};base64,{image_base64}"
                        }
                
                # Расширенный style_map для обработки стилей Word
                style_map = """
                p[style-name='Intense Emphasis'] => em
                p[style-name='emphasis'] => em
                r[style-name='Intense Emphasis'] => em
                r[style-name='emphasis'] => em
                p[style-name='Strong'] => strong
                r[style-name='Strong'] => strong
                p[style-name='Heading 1'] => h1
                p[style-name='Heading 2'] => h2
                p[style-name='Heading 3'] => h3
                p[style-name='Heading 4'] => h4
                p[style-name='Heading 5'] => h5
                p[style-name='Heading 6'] => h6
                p[style-name='Заголовок 1'] => h1
                p[style-name='Заголовок 2'] => h2
                p[style-name='Заголовок 3'] => h3
                p[style-name='Заголовок 4'] => h4
                p[style-name='Заголовок 5'] => h5
                p[style-name='Заголовок 6'] => h6
                p[style-name='Title'] => h1
                p[style-name='Subtitle'] => h2
                p[style-name='List Paragraph'] => p
                p[style-name='Quote'] => blockquote
                p[style-name='Caption'] => p.caption
                p[style-name='Table Caption'] => p.table-caption
                """
                
                # Используем mammoth для конвертации с обработкой изображений и стилей
                result = mammoth.convert_to_html(
                    word_file, 
                    convert_image=mammoth.images.img_element(convert_image),
                    style_map=style_map
                )
                html_content = result.value
                mammoth_messages = result.messages
                
                converter_used = 'mammoth'
                
                if mammoth_messages:
                    warnings.extend([str(m) for m in mammoth_messages])
                    
            except Exception as mammoth_error:
                return Response(
                    {'error': f'Ошибка при обработке файла: {str(mammoth_error)}'}, 
                    status=status.HTTP_400_BAD_REQUEST
                )
        
        # Постобработка HTML: сохранение стилей и конвертация в Bootstrap
        try:
            processor = WordImportProcessor()
            html_content = processor.process_html(html_content)
        except Exception as processing_error:
            # Если постобработка не удалась, возвращаем исходный контент
            warnings.append(f'Предупреждение: ошибка при постобработке HTML: {str(processing_error)}')
        
        return Response({
            'content': html_content,
            'warnings': warnings,
            'converter_used': converter_used
        })


class ArticleVersionViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра версий статей (только для пользователей с полными правами на категорию)
    """
    queryset = ArticleVersion.objects.all()
    serializer_class = ArticleVersionSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        queryset = ArticleVersion.objects.all()
        article_id = self.request.query_params.get('article', None)
        if article_id:
            queryset = queryset.filter(article_id=article_id)
        
        # Фильтруем по правам доступа
        user = self.request.user
        
        # Суперпользователи видят все версии
        if user.is_superuser:
            return queryset.select_related('article', 'author')
        
        # Фильтруем версии по системным правам
        from .permissions import ArticlePermission
        permission_checker = ArticlePermission()
        
        # Получаем итоговые системные права пользователя (группы + is_staff)
        system_permission = permission_checker._get_effective_system_permission(user)
        
        # Только пользователи с правами редактирования (edit) имеют доступ к версиям
        if system_permission == 'edit':
            return queryset.select_related('article', 'author')
        
        # Пользователи с read или без прав - показываем только версии своих статей
        allowed_articles = []
        for version in queryset.select_related('article', 'author'):
            article = version.article
            # Автор статьи имеет доступ к версиям
            if article.author == user:
                allowed_articles.append(version.id)
        
        return queryset.filter(id__in=allowed_articles).select_related('article', 'author')


class ArticleImageViewSet(viewsets.ModelViewSet):
    """
    ViewSet для управления изображениями статей
    """
    queryset = ArticleImage.objects.all()
    serializer_class = ArticleImageSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        queryset = ArticleImage.objects.all()
        article_id = self.request.query_params.get('article', None)
        if article_id:
            queryset = queryset.filter(article_id=article_id)
        return queryset.select_related('article', 'uploaded_by')
    
    def perform_create(self, serializer):
        serializer.save(uploaded_by=self.request.user)


class TechnologyViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра технологий (только чтение, создание/редактирование через админку)
    Показывает все технологии (права проверяются на уровне статей)
    """
    serializer_class = TechnologySerializer
    permission_classes = [permissions.IsAuthenticated]  # Требуется авторизация
    pagination_class = None  # Отключаем пагинацию для технологий
    
    def get_queryset(self):
        """Возвращает все технологии"""
        return Technology.objects.all().order_by('sort_order', 'name').prefetch_related('categories')
    
    def get_serializer_context(self):
        """Передаем request в контекст сериализатора для фильтрации категорий"""
        context = super().get_serializer_context()
        context['request'] = self.request
        return context


class CategoryViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра категорий (только чтение, создание/редактирование через админку)
    Показывает все категории (права проверяются на уровне статей)
    """
    serializer_class = CategorySerializer
    permission_classes = [permissions.IsAuthenticated]  # Требуется авторизация
    pagination_class = None  # Отключаем пагинацию для категорий
    
    def get_queryset(self):
        """Возвращает все категории"""
        return Category.objects.all().order_by('technology', 'sort_order', 'name').select_related('technology').prefetch_related('models')
    
    def get_serializer_context(self):
        """Передает request в контекст сериализатора для определения прав пользователя"""
        context = super().get_serializer_context()
        context['request'] = self.request
        return context


class ModelViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра моделей (только чтение, создание/редактирование через админку)
    Показывает все модели (права проверяются на уровне статей)
    """
    serializer_class = ModelSerializer
    permission_classes = [permissions.IsAuthenticated]  # Требуется авторизация
    pagination_class = None  # Отключаем пагинацию для моделей
    
    def get_queryset(self):
        """Возвращает все модели с возможностью фильтрации по категории"""
        queryset = Model.objects.all().order_by('category', 'sort_order', 'name').select_related('category')
        
        # Фильтр по категории
        category_id = self.request.query_params.get('category', None)
        if category_id:
            queryset = queryset.filter(category_id=category_id)
        
        return queryset
    
    def get_serializer_context(self):
        """Передает request в контекст сериализатора для определения прав пользователя"""
        context = super().get_serializer_context()
        context['request'] = self.request
        return context


class TagViewSet(viewsets.ModelViewSet):
    """
    ViewSet для управления тегами
    """
    queryset = Tag.objects.all().order_by('name')
    serializer_class = TagSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_queryset(self):
        queryset = Tag.objects.all().order_by('name')
        # Поиск по названию (регистронезависимый для SQLite)
        search = self.request.query_params.get('search', None)
        if search:
            # В SQLite LOWER() не работает с кириллицей, используем COLLATE NOCASE
            # или фильтруем через Python. Используем __iregex для регистронезависимого поиска
            import re
            search_escaped = re.escape(search)
            queryset = queryset.filter(name__iregex=search_escaped)
        return queryset
    
    def create(self, request, *args, **kwargs):
        # Проверяем, существует ли уже тег с таким именем (без учета регистра)
        name = request.data.get('name', '').strip()
        if name:
            # Проверяем по имени (без учета регистра)
            existing_tag = Tag.objects.filter(name__iexact=name).first()
            if existing_tag:
                # Если тег существует, возвращаем его вместо создания нового
                serializer = self.get_serializer(existing_tag)
                return Response(serializer.data, status=status.HTTP_200_OK)
            
            # Также проверяем по slug (на случай, если имя отличается, но slug одинаковый)
            from django.utils.text import slugify
            slug = slugify(name)
            existing_tag_by_slug = Tag.objects.filter(slug=slug).first()
            if existing_tag_by_slug:
                # Если тег с таким slug существует, возвращаем его
                serializer = self.get_serializer(existing_tag_by_slug)
                return Response(serializer.data, status=status.HTTP_200_OK)
        
        # Если тег не найден, создаем новый через стандартный метод
        return super().create(request, *args, **kwargs)


class ArticleOptionViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра опций статей (только чтение, создание/редактирование через админку)
    """
    queryset = ArticleOption.objects.all().order_by('sort_order', 'name')
    serializer_class = ArticleOptionSerializer
    permission_classes = [permissions.AllowAny]  # Опции доступны всем для выбора


class GroupViewSet(viewsets.ModelViewSet):
    """
    ViewSet для управления группами пользователей
    """
    queryset = Group.objects.all().prefetch_related('users')
    permission_classes = [permissions.IsAuthenticated]
    
    def get_serializer_class(self):
        if self.action == 'retrieve':
            return GroupDetailSerializer
        return GroupSerializer
    
    def get_queryset(self):
        queryset = Group.objects.all().prefetch_related('users')
        
        # Поиск по названию
        search = self.request.query_params.get('search', None)
        if search:
            queryset = queryset.filter(name__icontains=search)
        
        # Фильтр по пользователю
        user_id = self.request.query_params.get('user', None)
        if user_id:
            queryset = queryset.filter(users__id=user_id).distinct()
        
        return queryset.order_by('name')


class ArticleTemplateViewSet(viewsets.ReadOnlyModelViewSet):
    """
    ViewSet для просмотра шаблонов статей (только чтение, создание/редактирование через админку)
    """
    queryset = ArticleTemplate.objects.all().order_by('name')
    serializer_class = ArticleTemplateSerializer
    permission_classes = [permissions.IsAuthenticated]  # Требуется авторизация


class CommentViewSet(viewsets.ModelViewSet):
    """
    ViewSet для управления комментариями к статьям
    Комментировать можно только опубликованные статьи
    Доступ имеют пользователи с правами read/edit
    Удаление комментариев отключено
    """
    queryset = Comment.objects.all()
    serializer_class = CommentSerializer
    permission_classes = [ArticlePermission]
    http_method_names = ['get', 'post', 'put', 'patch', 'head', 'options']  # Исключаем delete
    
    def get_queryset(self):
        """Фильтруем комментарии по статье"""
        # Для операций list (GET /comments/) возвращаем только родительские комментарии
        # Для операций retrieve, update, destroy (GET/PUT/PATCH/DELETE /comments/{id}/) возвращаем все комментарии
        if self.action == 'list':
            queryset = Comment.objects.filter(parent__isnull=True).select_related('author', 'article', 'parent').prefetch_related('referenced_comments', 'replies')
            article_id = self.request.query_params.get('article', None)
            if article_id:
                queryset = queryset.filter(article_id=article_id)
            return queryset.order_by('created_at')
        else:
            # Для других операций (retrieve, update, destroy) возвращаем все комментарии
            queryset = Comment.objects.all().select_related('author', 'article', 'parent').prefetch_related('referenced_comments', 'replies')
            return queryset
    
    def perform_create(self, serializer):
        """Создание комментария с проверкой прав"""
        article = serializer.validated_data['article']
        
        # Проверяем, что статья опубликована
        if not article.is_published:
            raise ValidationError({
                'article': 'Комментировать можно только опубликованные статьи'
            })
        
        # Проверяем права доступа через ArticlePermission
        from .permissions import ArticlePermission
        permission_checker = ArticlePermission()
        
        user = self.request.user
        if not user.is_superuser:
            system_permission = permission_checker._get_effective_system_permission(user)
            if system_permission not in ['read', 'edit']:
                raise permissions.PermissionDenied('У вас нет прав на комментирование статей')
        
        serializer.save(author=user)
    
    def perform_update(self, serializer):
        """Обновление комментария - только автор может редактировать"""
        comment = self.get_object()
        user = self.request.user
        
        # Только автор может редактировать свой комментарий
        if comment.author != user and not user.is_superuser:
            raise permissions.PermissionDenied('Вы можете редактировать только свои комментарии')
        
        serializer.save()

